#!/usr/bin/env python3
"""
Export API for LibreChat - Provides PDF/DOCX conversion endpoints
Pure Python implementation without external scripts
"""

from http.server import HTTPServer, BaseHTTPRequestHandler
import json
import hashlib
import time
import tempfile
import os
import io

class ExportAPIHandler(BaseHTTPRequestHandler):
    GOTENBERG_URL = os.getenv('GOTENBERG_URL', 'http://gotenberg:3000')

    def do_OPTIONS(self):
        """Handle CORS preflight"""
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, GET, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()

    def markdown_to_html(self, markdown, title="Conversation"):
        """Convert markdown to HTML with basic styling"""
        # Escape HTML characters first
        html = markdown.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

        # Convert newlines to <br> temporarily
        lines = html.split('\n')
        formatted_lines = []

        for line in lines:
            # Headers
            if line.startswith('#### '):
                formatted_lines.append(f'<h4>{line[5:]}</h4>')
            elif line.startswith('### '):
                formatted_lines.append(f'<h3>{line[4:]}</h3>')
            elif line.startswith('## '):
                formatted_lines.append(f'<h2>{line[3:]}</h2>')
            elif line.startswith('# '):
                formatted_lines.append(f'<h1>{line[2:]}</h1>')
            elif line.strip() == '---':
                formatted_lines.append('<hr>')
            elif line.strip():
                formatted_lines.append(f'<p>{line}</p>')
            else:
                formatted_lines.append('<br>')

        html = '\n'.join(formatted_lines)

        return f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
            max-width: 800px;
            margin: 40px auto;
            padding: 20px;
            line-height: 1.6;
            color: #333;
        }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #34495e; margin-top: 24px; }}
        h3 {{ color: #7f8c8d; }}
        pre {{
            background: #f4f4f4;
            padding: 12px;
            border-radius: 5px;
            overflow-x: auto;
            border-left: 3px solid #3498db;
        }}
        code {{
            background: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
        }}
        p {{ margin: 12px 0; }}
        hr {{ border: none; border-top: 1px solid #ddd; margin: 24px 0; }}
    </style>
</head>
<body>
{html}
</body>
</html>'''

    def convert_to_pdf(self, html_content):
        """Convert HTML to PDF using Gotenberg"""
        try:
            import urllib.request

            # Create multipart form data
            boundary = '----WebKitFormBoundary' + hashlib.md5(str(time.time()).encode()).hexdigest()[:16]

            body = []
            body.append(f'--{boundary}'.encode())
            body.append(b'Content-Disposition: form-data; name="files"; filename="index.html"')
            body.append(b'Content-Type: text/html')
            body.append(b'')
            body.append(html_content.encode('utf-8'))
            body.append(f'--{boundary}--'.encode())
            body.append(b'')

            body_bytes = b'\r\n'.join(body)

            # Send request to Gotenberg
            url = f'{self.GOTENBERG_URL}/forms/chromium/convert/html'
            req = urllib.request.Request(url, data=body_bytes)
            req.add_header('Content-Type', f'multipart/form-data; boundary={boundary}')

            with urllib.request.urlopen(req, timeout=30) as response:
                return response.read()

        except Exception as e:
            raise Exception(f'Gotenberg conversion failed: {str(e)}')

    def convert_to_docx(self, markdown):
        """Convert markdown to DOCX using basic formatting"""
        try:
            # Check if python-docx is available
            from docx import Document
            from docx.shared import Pt, RGBColor

            doc = Document()

            # Add title
            title = doc.add_heading('Conversation Export', 0)

            # Parse markdown and add to document
            lines = markdown.split('\n')
            current_paragraph = None

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('---'):
                    doc.add_paragraph('_' * 50)
                else:
                    doc.add_paragraph(line)

            # Save to bytes
            docx_io = io.BytesIO()
            doc.save(docx_io)
            docx_io.seek(0)
            return docx_io.read()

        except ImportError:
            # Fallback: return error message
            raise Exception('python-docx not installed. PDF export only.')

    def do_POST(self):
        """Convert markdown to PDF/DOCX"""
        if self.path == '/api/convert':
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)

            try:
                # Parse JSON with better error handling
                try:
                    data = json.loads(post_data.decode('utf-8'))
                except json.JSONDecodeError as e:
                    print(f'❌ JSON decode error: {e}')
                    print(f'   Raw data: {post_data[:200]}')
                    raise Exception(f'Invalid JSON: {str(e)}')

                markdown = data.get('markdown', '')
                format_type = data.get('format', 'pdf')

                if not markdown:
                    self.send_error(400, 'No markdown content provided')
                    return

                # Create unique filename
                timestamp = int(time.time())
                hash_id = hashlib.md5(markdown.encode()).hexdigest()[:8]
                filename = f'conversation-{timestamp}-{hash_id}.{format_type}'

                # Convert based on format
                if format_type == 'pdf':
                    html_content = self.markdown_to_html(markdown, 'Conversation Export')
                    file_content = self.convert_to_pdf(html_content)
                    content_type = 'application/pdf'
                elif format_type == 'docx':
                    file_content = self.convert_to_docx(markdown)
                    content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                else:
                    self.send_error(400, f'Unsupported format: {format_type}')
                    return

                # Send file
                self.send_response(200)
                self.send_header('Access-Control-Allow-Origin', '*')
                self.send_header('Content-Type', content_type)
                self.send_header('Content-Disposition', f'attachment; filename="{filename}"')
                self.send_header('Content-Length', str(len(file_content)))
                self.end_headers()
                self.wfile.write(file_content)

                print(f'✅ Converted and sent {filename}')

            except Exception as e:
                print(f'❌ Conversion error: {str(e)}')
                self.send_error(500, str(e))
        else:
            self.send_error(404)

    def do_GET(self):
        """Health check"""
        if self.path == '/health':
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(json.dumps({'status': 'ok'}).encode())
        else:
            self.send_error(404)

def run(port=8081):
    server_address = ('0.0.0.0', port)
    httpd = HTTPServer(server_address, ExportAPIHandler)
    print(f'🚀 Export API v2 running on port {port}')
    print(f'📝 Gotenberg URL: {ExportAPIHandler.GOTENBERG_URL}')
    print(f'✅ Ready to serve conversion requests')
    httpd.serve_forever()

if __name__ == '__main__':
    run()
